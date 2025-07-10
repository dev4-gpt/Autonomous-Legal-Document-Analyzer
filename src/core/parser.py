"""
Enhanced document parser for the Legal Document Analyzer.
Supports multiple file formats with robust error handling and performance optimization.
"""

import os
import time
from pathlib import Path
from typing import Tuple, Optional, Dict, Any
from dataclasses import dataclass
import hashlib

# Document processing imports
import fitz  # PyMuPDF
import docx
from bs4 import BeautifulSoup
import magic  # python-magic for file type detection

from src.config import config
from src.utils import logger, performance_logger
from src.database import Document, with_db_session


@dataclass
class ParseResult:
    """Result of document parsing operation."""
    success: bool
    text: str
    doc_id: str
    metadata: Dict[str, Any]
    error_message: Optional[str] = None
    processing_time: float = 0.0


class DocumentParser:
    """Enhanced document parser with comprehensive format support."""
    
    def __init__(self):
        self.supported_extensions = config.SUPPORTED_EXTENSIONS
        self.max_file_size = config.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
    
    def parse_file(self, file_path: str) -> ParseResult:
        """
        Parse a document file and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParseResult with extracted text and metadata
        """
        start_time = time.time()
        file_path = Path(file_path)
        
        try:
            # Validate file
            validation_result = self._validate_file(file_path)
            if not validation_result["valid"]:
                return ParseResult(
                    success=False,
                    text="",
                    doc_id="",
                    metadata={},
                    error_message=validation_result["error"],
                    processing_time=time.time() - start_time
                )
            
            # Generate document ID
            doc_id = self._generate_doc_id(file_path)
            
            # Get file metadata
            metadata = self._get_file_metadata(file_path)
            
            # Parse based on file extension
            ext = file_path.suffix.lower()
            if ext == '.pdf':
                text = self._parse_pdf(file_path)
            elif ext == '.docx':
                text = self._parse_docx(file_path)
            elif ext == '.txt':
                text = self._parse_txt(file_path)
            elif ext == '.html':
                text = self._parse_html(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
            
            # Clean and validate extracted text
            text = self._clean_text(text)
            if not text.strip():
                raise ValueError("No text content extracted from document")
            
            processing_time = time.time() - start_time
            
            # Log performance metrics
            performance_logger.log_timing(
                "document_parsing",
                processing_time,
                file_size=metadata["file_size"],
                file_type=ext,
                text_length=len(text)
            )
            
            logger.info(f"Successfully parsed {file_path.name} ({len(text)} characters)")
            
            return ParseResult(
                success=True,
                text=text,
                doc_id=doc_id,
                metadata=metadata,
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"Failed to parse {file_path.name}: {str(e)}"
            logger.error(error_msg, exception=e)
            
            return ParseResult(
                success=False,
                text="",
                doc_id=self._generate_doc_id(file_path) if file_path.exists() else "",
                metadata=self._get_file_metadata(file_path) if file_path.exists() else {},
                error_message=error_msg,
                processing_time=processing_time
            )
    
    def _validate_file(self, file_path: Path) -> Dict[str, Any]:
        """Validate file before processing."""
        if not file_path.exists():
            return {"valid": False, "error": "File does not exist"}
        
        if not file_path.is_file():
            return {"valid": False, "error": "Path is not a file"}
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            return {
                "valid": False, 
                "error": f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds limit ({config.MAX_FILE_SIZE_MB}MB)"
            }
        
        if file_size == 0:
            return {"valid": False, "error": "File is empty"}
        
        # Check file extension
        ext = file_path.suffix.lower()
        if ext not in self.supported_extensions:
            return {
                "valid": False, 
                "error": f"Unsupported file type: {ext}. Supported: {', '.join(self.supported_extensions)}"
            }
        
        # Verify file type using magic numbers (more reliable than extension)
        try:
            mime_type = magic.from_file(str(file_path), mime=True)
            expected_mimes = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.txt': 'text/plain',
                '.html': 'text/html'
            }
            
            if ext in expected_mimes and not mime_type.startswith(expected_mimes[ext].split('/')[0]):
                logger.warning(f"File extension {ext} doesn't match MIME type {mime_type}")
        except Exception as e:
            logger.warning(f"Could not verify MIME type for {file_path}: {e}")
        
        return {"valid": True, "error": None}
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID based on file content and metadata."""
        try:
            # Use file content hash + filename + size for uniqueness
            with open(file_path, 'rb') as f:
                content_hash = hashlib.md5(f.read()).hexdigest()[:16]
            
            file_info = f"{file_path.name}_{file_path.stat().st_size}_{content_hash}"
            return hashlib.sha256(file_info.encode()).hexdigest()[:32]
        except Exception:
            # Fallback to filename-based ID
            return hashlib.sha256(str(file_path).encode()).hexdigest()[:32]
    
    def _get_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract file metadata."""
        try:
            stat = file_path.stat()
            return {
                "filename": file_path.name,
                "file_path": str(file_path),
                "file_size": stat.st_size,
                "file_type": file_path.suffix.lower(),
                "created_time": stat.st_ctime,
                "modified_time": stat.st_mtime,
                "mime_type": magic.from_file(str(file_path), mime=True) if magic else None
            }
        except Exception as e:
            logger.warning(f"Could not extract metadata for {file_path}: {e}")
            return {
                "filename": file_path.name,
                "file_path": str(file_path),
                "file_size": 0,
                "file_type": file_path.suffix.lower(),
                "created_time": 0,
                "modified_time": 0,
                "mime_type": None
            }
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF document using PyMuPDF."""
        try:
            doc = fitz.open(str(file_path))
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            doc.close()
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise ValueError(f"PDF parsing failed: {str(e)}")
    
    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX document using python-docx."""
        try:
            doc = docx.Document(str(file_path))
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise ValueError(f"DOCX parsing failed: {str(e)}")
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse plain text file with encoding detection."""
        try:
            # Try common encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors ignored
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            raise ValueError(f"Text file parsing failed: {str(e)}")
    
    def _parse_html(self, file_path: Path) -> str:
        """Parse HTML document using BeautifulSoup."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text and clean up whitespace
            text = soup.get_text(separator='\n')
            return text
            
        except Exception as e:
            raise ValueError(f"HTML parsing failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Normalize whitespace
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            if line:
                lines.append(line)
        
        # Join lines and normalize spaces
        cleaned = '\n'.join(lines)
        
        # Remove excessive whitespace
        import re
        cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned)  # Max 2 consecutive newlines
        cleaned = re.sub(r' +', ' ', cleaned)  # Multiple spaces to single space
        
        return cleaned.strip()


# Global parser instance
document_parser = DocumentParser()

# Convenience function
def parse_document(file_path: str) -> ParseResult:
    """Parse a document file and return the result."""
    return document_parser.parse_file(file_path)

@with_db_session
def parse_and_store_document(session, file_path: str) -> Tuple[bool, str]:
    """Parse document and store metadata in database."""
    result = parse_document(file_path)
    
    try:
        # Create document record
        document = Document(
            doc_id=result.doc_id,
            filename=result.metadata.get("filename", ""),
            file_path=result.metadata.get("file_path", ""),
            file_size=result.metadata.get("file_size", 0),
            file_type=result.metadata.get("file_type", ""),
            processing_status="completed" if result.success else "failed",
            processing_time=result.processing_time,
            error_message=result.error_message
        )
        
        session.add(document)
        session.commit()
        
        return result.success, result.doc_id if result.success else result.error_message
        
    except Exception as e:
        logger.error(f"Failed to store document metadata: {e}")
        return False, str(e)
