"""
Pytest configuration and fixtures for the Legal Document Analyzer tests.
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import Mock, patch
import sys

# Add src to path for testing
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from src.config import TestingConfig
from src.database import db_manager, init_database
from src.database.models import Base


@pytest.fixture(scope="session")
def test_config():
    """Provide test configuration."""
    return TestingConfig()


@pytest.fixture(scope="function")
def temp_dir():
    """Create a temporary directory for test files."""
    temp_dir = tempfile.mkdtemp()
    yield Path(temp_dir)
    shutil.rmtree(temp_dir)


@pytest.fixture(scope="function")
def test_database():
    """Create a test database."""
    # Use in-memory SQLite for testing
    with patch('src.config.config.DATABASE_URL', 'sqlite:///:memory:'):
        # Reinitialize database manager with test config
        db_manager._initialize_database()
        db_manager.create_tables()
        yield db_manager
        # Cleanup is automatic with in-memory database


@pytest.fixture(scope="function")
def db_session(test_database):
    """Provide a database session for testing."""
    with test_database.get_session() as session:
        yield session


@pytest.fixture
def sample_contract_text():
    """Provide sample contract text for testing."""
    return """
    CONFIDENTIALITY AGREEMENT
    
    This Confidentiality Agreement ("Agreement") is entered into on [DATE] between 
    Company A ("Disclosing Party") and Company B ("Receiving Party").
    
    1. CONFIDENTIAL INFORMATION
    The Receiving Party acknowledges that it may receive confidential information
    from the Disclosing Party.
    
    2. OBLIGATIONS
    The Receiving Party agrees to:
    a) Keep all confidential information strictly confidential
    b) Not disclose to any third parties
    c) Use information only for the intended purpose
    
    3. TERM
    This Agreement shall remain in effect for a period of 5 years.
    
    4. TERMINATION
    Either party may terminate this Agreement with 30 days written notice.
    
    5. GOVERNING LAW
    This Agreement shall be governed by the laws of [STATE].
    """


@pytest.fixture
def sample_pdf_file(temp_dir):
    """Create a sample PDF file for testing."""
    pdf_path = temp_dir / "test_contract.pdf"
    
    # Create a simple PDF using reportlab if available, otherwise create a dummy file
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test Contract Document")
        c.drawString(100, 700, "This is a test contract for unit testing.")
        c.save()
    except ImportError:
        # Fallback: create a dummy file
        pdf_path.write_bytes(b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n")
    
    return pdf_path


@pytest.fixture
def sample_docx_file(temp_dir):
    """Create a sample DOCX file for testing."""
    docx_path = temp_dir / "test_contract.docx"
    
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Contract', 0)
        doc.add_paragraph('This is a test contract document for unit testing.')
        doc.add_paragraph('It contains sample legal text for analysis.')
        doc.save(str(docx_path))
    except ImportError:
        # Fallback: create a dummy file
        docx_path.write_bytes(b"PK\x03\x04")  # ZIP file signature
    
    return docx_path


@pytest.fixture
def sample_txt_file(temp_dir, sample_contract_text):
    """Create a sample text file for testing."""
    txt_path = temp_dir / "test_contract.txt"
    txt_path.write_text(sample_contract_text)
    return txt_path


@pytest.fixture
def mock_llm_response():
    """Mock LLM response for testing."""
    def _mock_response(content="Test response", error=None):
        mock = Mock()
        mock.content = content
        mock.provider = "test"
        mock.model = "test-model"
        mock.processing_time = 0.1
        mock.token_count = 10
        mock.error = error
        return mock
    return _mock_response


@pytest.fixture
def mock_llm_manager(mock_llm_response):
    """Mock LLM manager for testing."""
    with patch('src.core.llm_manager.llm_manager') as mock_manager:
        mock_manager.invoke.return_value = mock_llm_response()
        mock_manager.current_provider = "test"
        mock_manager.providers = {
            "test": {
                "client": Mock(),
                "model": "test-model",
                "available": True
            }
        }
        yield mock_manager


@pytest.fixture
def mock_document_parser():
    """Mock document parser for testing."""
    with patch('src.core.parser.document_parser') as mock_parser:
        mock_result = Mock()
        mock_result.success = True
        mock_result.text = "Sample contract text"
        mock_result.doc_id = "test_doc_123"
        mock_result.metadata = {
            "filename": "test.pdf",
            "file_size": 1024,
            "file_type": ".pdf"
        }
        mock_result.processing_time = 0.5
        mock_result.error_message = None
        
        mock_parser.parse_file.return_value = mock_result
        yield mock_parser


@pytest.fixture(autouse=True)
def setup_test_environment():
    """Set up test environment before each test."""
    # Ensure test directories exist
    test_dirs = [
        Path("data/test"),
        Path("logs/test")
    ]
    
    for dir_path in test_dirs:
        dir_path.mkdir(parents=True, exist_ok=True)
    
    yield
    
    # Cleanup after test
    for dir_path in test_dirs:
        if dir_path.exists():
            shutil.rmtree(dir_path, ignore_errors=True)
